import streamlit as st
import re, os, json, tempfile
from io import BytesIO
from datetime import datetime
from docx import Document
from pdfminer.high_level import extract_text
import openai

# ─── CONFIG ────────────────────────────────────────────────
openai.api_key = st.secrets["OPENAI_API_KEY"]
OPENAI_MODEL = "gpt-4o"

# ─── FORMAT DATE ───────────────────────────────────────────
def format_certificate_date(raw_date_str):
    try:
        dt = datetime.strptime(raw_date_str, "%B %d, %Y")
    except ValueError:
        for fmt in ("%m/%d/%Y", "%Y-%m-%d"):
            try:
                dt = datetime.strptime(raw_date_str, fmt)
                break
            except ValueError:
                continue
        else:
            return "Dated ______"
    day = dt.day
    suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    month = dt.strftime("%B")
    year_words = {
        "2025": "Two Thousand and Twenty-Five",
        "2024": "Two Thousand and Twenty-Four",
        "2023": "Two Thousand and Twenty-Three"
    }.get(dt.strftime("%Y"), dt.strftime("%Y"))
    return f"Dated the {day}{suffix} of {month}\n{year_words}"

# ─── FALLBACK MESSAGE ──────────────────────────────────────
def fallback_commendation(name, title, org):
    title_lower = title.lower()
    if "award" in title_lower or "of the year" in title_lower or "honoree" in title_lower:
        message = f"On behalf of the California State Legislature, congratulations on being recognized as {org}'s {title}. "
    elif any(kw in title_lower for kw in ["president", "board", "officer", "service", "chair", "director"]):
        message = f"On behalf of the California State Legislature, thank you for your service as {title} with {org}. "
    elif "opening" in title_lower or "grand" in title_lower:
        message = f"On behalf of the California State Legislature, congratulations on the opening of {org}. "
    elif "graduat" in title_lower:
        message = f"On behalf of the California State Legislature, congratulations on successfully graduating from {org}. "
    else:
        message = f"On behalf of the California State Legislature, we commend you for your accomplishments with {org}. "
    message += "This recognition speaks highly of your dedication and contributions to the community."
    return message

# ─── TONE CATEGORY ─────────────────────────────────────────
def categorize_tone(title):
    title_lower = title.lower()
    if any(kw in title_lower for kw in ["award", "of the year", "honoree", "achievement", "excellence", "inductee"]):
        return "🏆 Award"
    elif any(kw in title_lower for kw in ["president", "officer", "board", "director", "retire", "service", "chair"]):
        return "👥 Service"
    elif any(kw in title_lower for kw in ["opening", "grand", "event", "dedication", "launch"]):
        return "🏛 Event"
    elif any(kw in title_lower for kw in ["graduate", "class of", "commencement"]):
        return "🎓 Graduation"
    else:
        return "📝 Recognition"

# ─── UI ────────────────────────────────────────────────────
st.set_page_config(layout="centered")
st.title("📄 Certificate Generator (Multi-Entry PDF)")
st.markdown("Upload a multi-entry recognition request PDF. This tool generates GPT-powered commendations and a multi-page Word document.")

pdf_file = st.file_uploader("📎 Upload Multi-Request PDF", type=["pdf"])
if not pdf_file:
    st.stop()

# ─── EXTRACT TEXT ──────────────────────────────────────────
with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
    tmp.write(pdf_file.read())
    tmp_path = tmp.name
pdf_text = extract_text(tmp_path)
os.remove(tmp_path)

entries = re.split(r"\_{5,}[\s\S]+?Stan Ellis\s+Assemblyman,? \d{1,2}(st|nd|rd|th)? District", pdf_text)
entries = [e.strip() for e in entries if e.strip()]
st.info(f"📄 {len(entries)} entries detected.")

# ─── GPT PROCESSING ────────────────────────────────────────
certs = []
for idx, entry in enumerate(entries):
    with st.spinner(f"🔍 Processing entry #{idx+1}..."):
        try:
            response = openai.ChatCompletion.create(
                model=OPENAI_MODEL,
                messages=[
                    {"role": "system", "content":
                     "You will be given a recognition request. Extract:\n"
                     "• name\n• title (award or position)\n• organization\n• date_raw\n\n"
                     "Then, generate a formal commendation message that:\n"
                     "- Begins with 'On behalf of the California State Legislature,'\n"
                     "- If an award: congratulates them on being recognized as OR receiving [org]'s [title]\n"
                     "- If service: thanks them for their role\n"
                     "- If a grand opening: congratulates on the opening of [org]\n"
                     "Conclude with 1–2 sentences affirming their impact.\n\n"
                     "Respond ONLY in valid JSON with keys: name, title, organization, date_raw, commendation."},
                    {"role": "user", "content": entry}
                ],
                temperature=0
            )
            raw = response["choices"][0]["message"]["content"].strip().strip("```json").strip("```")
            parsed = json.loads(raw)
            parsed["formatted_date"] = format_certificate_date(parsed["date_raw"])
            parsed["tone"] = categorize_tone(parsed["title"])
            certs.append(parsed)
        except Exception:
            fallback = {
                "name": "UNKNOWN",
                "title": "UNKNOWN",
                "organization": "UNKNOWN",
                "date_raw": "UNKNOWN",
                "formatted_date": "Dated ______",
                "commendation": fallback_commendation("UNKNOWN", "UNKNOWN", "UNKNOWN"),
                "tone": "📝 Recognition"
            }
            certs.append(fallback)

if not certs:
    st.error("❌ No valid entries parsed.")
    st.stop()

# ─── PREVIEW UI ────────────────────────────────────────────
st.subheader("👁 Certificate Preview")
for i, cert in enumerate(certs, 1):
    with st.expander(f"{cert['tone']} #{i}: {cert['name']} – {cert['title']}"):
        st.write(f"**Organization:** {cert['organization']}")
        st.write(f"**Date:** {cert['formatted_date']}")
        st.text_area("Commendation Message", cert["commendation"], height=100)

# ─── GENERATE .DOCX ────────────────────────────────────────
if st.button("📄 Generate Word Document"):
    template_path = "cert_template.docx"
    if not os.path.exists(template_path):
        st.error("❌ Missing 'cert_template.docx' in app folder.")
        st.stop()

    merged_doc = Document()
    for idx, cert in enumerate(certs):
        doc = Document(template_path)
        for p in doc.paragraphs:
            for key, val in {
                "«Name»": cert["name"],
                "«Title»": cert["title"],
                "«Certificate_Text»": cert["commendation"],
                "«Formatted_Date»": cert["formatted_date"]
            }.items():
                if key in p.text:
                    for run in p.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, val)
        for element in doc.element.body:
            merged_doc.element.body.append(element)
        if idx < len(certs) - 1:
            merged_doc.add_page_break()

    buffer = BytesIO()
    merged_doc.save(buffer)
    buffer.seek(0)

    st.success("🎉 Certificate document ready!")
    st.download_button(
        label="📥 Download Certificates (.docx)",
        data=buffer,
        file_name="Certificates_Merged.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
